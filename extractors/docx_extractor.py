

# -*- coding: utf-8 -*-
"""
extractors/docx_extractor.py

DOCX extractor:
- Reads .docx files (paragraphs + table cells)
- Normalizes whitespace and filters tiny fragments
- Emits LangChain Document objects with rich metadata

Outputs one Document per .docx file (by default), or optionally per-section
if you want to split into smaller units before your RAG chunker runs.

Metadata fields (aligned with PDF extractor & your RAG pipeline):
    file_name, source, page_number=1, doc_type, chunk_kind="text", created_at

Usage:
    from extractors.docx_extractor import DocxExtractor
    dx = DocxExtractor(min_text_len=20, doc_type_infer=True, join_tables=True)
    docs = dx.load_docx("path/to/file.docx")
    # or
    docs = dx.load_path("path/to/folder")

Notes:
- Requires `python-docx` (imported as `docx`).
- We treat DOCX as a single logical "page" (page_number=1), since pagination is not defined in the file format.
- If you want finer-grained chunks, keep `per_file_document=True` here and let your RAG chunker
  (rag/chunker.py) split the large text into retrieval-sized chunks consistently across sources.

"""

from __future__ import annotations

import os
import glob
from typing import List, Optional


# LangChain Document
try:
    from langchain_core.documents import Document
except Exception:
    class Document:  # minimal shim if langchain is not present at import-time
        def __init__(self, page_content: str, metadata: dict):
            self.page_content = page_content
            self.metadata = metadata

# python-docx
try:
    from docx import Document as DocxDocument  # type: ignore
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

# Shared helpers
try:
    from vector_store.base import normalize_text, now_iso
except Exception:
    import datetime
    def normalize_text(txt: Optional[str]) -> str:
        return " ".join((txt or "").split())
    def now_iso() -> str:
        return datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


class DocxExtractor:
    """
    Extracts text from DOCX files (paragraphs + tables) into LangChain Documents.

    Parameters
    ----------
    min_text_len : int
        Minimum normalized text length to accept a document (prevents blank outputs).
    doc_type_infer : bool
        Infer doc_type heuristically from file name (technical/schema/erd/functional/...).
    join_tables : bool
        If True, include table cell text (joined row-wise) with paragraphs.
    per_file_document : bool
        If True, return one Document per file (recommended — let the downstream chunker split).
        If False, returns multiple Documents (one per major section: paragraphs vs tables).
    """

    def __init__(
        self,
        min_text_len: int = 20,
        doc_type_infer: bool = True,
        join_tables: bool = True,
        per_file_document: bool = True,
    ) -> None:
        self.min_text_len = int(min_text_len)
        self.doc_type_infer = bool(doc_type_infer)
        self.join_tables = bool(join_tables)
        self.per_file_document = bool(per_file_document)

        if not _DOCX_AVAILABLE:
            print("[WARN] python-docx is not installed; DOCX extraction will be skipped.")

    # ----------------------------
    # Public API
    # ----------------------------

    def load_docx(self, path: str) -> List[Document]:
        """
        Extract text from a single DOCX file.
        """
        out: List[Document] = []
        file_name = os.path.basename(path)
        doc_type = self._infer_doc_type(file_name) if self.doc_type_infer else "unknown"

        if not _DOCX_AVAILABLE:
            return out

        try:
            d = DocxDocument(path)
        except Exception as e:
            print(f"[WARN] Cannot open DOCX {file_name}: {e}")
            return out

        # Collect paragraphs
        paragraphs: List[str] = []
        for p in d.paragraphs:
            t = normalize_text(getattr(p, "text", ""))
            if t:
                paragraphs.append(t)

        # Collect table text (optional)
        table_lines: List[str] = []
        if self.join_tables:
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [normalize_text(c.text) for c in row.cells if normalize_text(c.text)]
                    if cells:
                        table_lines.append(" | ".join(cells))

        # Emit
        base_meta = {
            "file_name": file_name,
            "source": path,
            "page_number": 1,               # DOCX has no stable pagination
            "doc_type": doc_type,
            "created_at": now_iso(),
        }

        if self.per_file_document:
            # One big doc, downstream chunker will split consistently across sources
            combined_parts: List[str] = []
            if paragraphs:
                combined_parts.append("\n".join(paragraphs))
            if table_lines:
                combined_parts.append("\n".join(table_lines))
            text = normalize_text("\n\n".join(combined_parts))
            if len(text) >= self.min_text_len:
                out.append(Document(page_content=text, metadata={**base_meta, "chunk_kind": "text"}))
        else:
            # Separate docs for paragraphs and tables (if present)
            if paragraphs:
                text = normalize_text("\n".join(paragraphs))
                if len(text) >= self.min_text_len:
                    out.append(Document(page_content=text, metadata={**base_meta, "chunk_kind": "text_paragraphs"}))
            if table_lines:
                text = normalize_text("\n".join(table_lines))
                if len(text) >= self.min_text_len:
                    out.append(Document(page_content=text, metadata={**base_meta, "chunk_kind": "text_tables"}))

        return out

    def load_path(self, path: str, *, recursive: bool = True) -> List[Document]:
        """
        Extract Documents from all DOCX files found under a folder (optionally recursive).
        """
        pattern = "**/*.docx" if recursive else "*.docx"
        files = sorted(glob.glob(os.path.join(path, pattern), recursive=recursive))
        all_docs: List[Document] = []
        for fp in files:
            all_docs.extend(self.load_docx(fp))
        if not all_docs:
            print(f"[INFO] No DOCX files processed under: {path}")
        return all_docs

    # ----------------------------
    # Internals
    # ----------------------------

    def _infer_doc_type(self, filename: str) -> str:
        """
        Heuristic doc_type inference based on file name tokens.
        """
        name = filename.lower()
        patterns = {
            "technical": ["tech", "technical", "architecture", "design"],
            "schema": ["schema", "tables", "columns", "ddl", "database"],
            "erd": ["erd", "entity", "diagram", "model"],
            "functional": ["functional", "business", "requirements", "brd"],
            "product": ["product", "features", "manual", "documentation", "guide"],
            "analytics": ["analytics", "kpi", "metric", "dashboard"],
            "reporting": ["report", "reporting"],
            "onboarding": ["onboarding", "getting_started", "introduction"],
            "user_guide": ["user_guide", "help", "how_to", "tutorial"],
            "api": ["api", "interface", "integration"],
        }
        for dt, keys in patterns.items():
            if any(k in name for k in keys):
                return dt
        return "unknown"
